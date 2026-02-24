import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls
import datetime
import re
from db_manager import DatabaseManager

# Константи
MONTHS_UA = {
    'січня': 1, 'лютого': 2, 'березня': 3, 'квітня': 4,
    'травня': 5, 'червня': 6, 'липня': 7, 'серпня': 8,
    'вересня': 9, 'жовтня': 10, 'листопада': 11, 'грудня': 12
}

class ReplacementForm(ttk.Frame):
    def __init__(self, parent, date_text=None, weekday=None, week_type=None, 
                 duty_group=None, duty_teacher=None, dorm_teacher=None,
                 replacements_dir=None, academic_year=None):
        super().__init__(parent)
        self.parent = parent
        
        # Зберігаємо дані з головної форми
        self.date_text = date_text
        self.weekday = weekday
        self.week_type = week_type
        self.duty_group_text = duty_group
        self.duty_teacher_text = duty_teacher
        self.dorm_teacher_text = dorm_teacher
        
        # Шлях для збереження файлів
        self.replacements_dir = replacements_dir
        self.academic_year = academic_year
        
        # Підключення до бази даних
        self.db = DatabaseManager()
        
        # Ледаче завантаження даних (завантажуються тільки при потребі)
        self._departments = None
        self._all_groups = None
        self._audiences = None
        self._disciplines = None
        self._groups = None
        
        # Тип практики (Виробнича або Переддипломна)
        self.practice_type = tk.StringVar(value="Виробнича")
        
        # Список замін буде заповнюватися пізніше
        
        # Список замін
        self.replacements = []
        
        self.create_widgets()
        
        # Ініціалізуємо відображення замін
        self.update_replacements_display()
    
    @property
    def departments(self):
        """Ледаче завантаження відділень"""
        if self._departments is None:
            self._departments = self.db.get_departments()
        return self._departments
    
    @property
    def all_groups(self):
        """Ледаче завантаження всіх груп"""
        if self._all_groups is None:
            self._all_groups = self.db.get_all_groups()
        return self._all_groups
    
    @property
    def audiences(self):
        """Ледаче завантаження аудиторій"""
        if self._audiences is None:
            self._audiences = self.db.get_audiences()
        return self._audiences
    
    @property
    def disciplines(self):
        """Ледаче завантаження дисциплін"""
        if self._disciplines is None:
            self._disciplines = self.db.get_disciplines()
        return self._disciplines
    
    @property
    def groups(self):
        """Ледаче завантаження груп за відділеннями"""
        if self._groups is None:
            self._groups = {}
            for dept in self.departments:
                self._groups[dept] = self.db.get_groups_by_department(dept)
        return self._groups
    
    def _clear_data_cache(self):
        """Очищення кешу даних"""
        self._departments = None
        self._all_groups = None
        self._audiences = None
        self._disciplines = None
        self._groups = None
    
    def create_widgets(self):
        # Створюємо Canvas та Scrollbar для прокручування всього вмісту
        canvas = tk.Canvas(self, bg="#f0f0f0")
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Додаємо підтримку прокручування колесиком миші для всього вікна
        def _on_canvas_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        # Прив'язуємо прокручування до canvas та всіх дочірніх елементів
        def bind_mousewheel_to_widget(widget):
            widget.bind("<MouseWheel>", _on_canvas_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel_to_widget(child)
        
        canvas.bind("<MouseWheel>", _on_canvas_mousewheel)
        
        # Зберігаємо посилання для подальшого використання
        self.canvas = canvas
        self.bind_mousewheel_to_widget = bind_mousewheel_to_widget
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Основний фрейм тепер всередині scrollable_frame
        main_frame = ttk.Frame(scrollable_frame, padding="8")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Створюємо заголовок вікна
        header_frame = ttk.Frame(main_frame, style='Header.TFrame')
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        header_label = ttk.Label(header_frame, text="Форма замін", font=("Arial", 12, "bold"), foreground="#ffffff", background="#4a6984")
        header_label.pack(fill=tk.X, ipady=8)
        
        # Видалено верхній радіогруп для вибору типу практики
        
        # Вибір відділення
        dept_frame = ttk.Frame(main_frame, style='TFrame')
        dept_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(dept_frame, text="Відділення:", font=("Arial", 10)).pack(side=tk.LEFT, padx=5)
        self.dept_var = tk.StringVar()
        self.dept_combo = ttk.Combobox(dept_frame, textvariable=self.dept_var, width=25, values=self.departments, font=("Arial", 10))
        self.dept_combo.pack(side=tk.LEFT, padx=5)
        self.dept_combo.bind("<<ComboboxSelected>>", self.on_department_change)
        
        # Фрейм для додавання замін
        replacement_frame = ttk.LabelFrame(main_frame, text="Додати заміну")
        replacement_frame.pack(fill=tk.X, pady=10)
        
        # Стилізуємо заголовок фрейму
        replacement_frame.configure(padding=10)
        
        # Змінюємо стиль заголовка фрейму
        style = ttk.Style()
        style.configure('TLabelframe.Label', font=("Arial", 12, "bold"))
        
        # Створюємо фрейм для полів вводу
        input_frame = ttk.Frame(replacement_frame)
        input_frame.pack(fill=tk.X, pady=10)
        
        # Налаштовуємо розміри стовпців для кращого розміщення
        input_frame.columnconfigure(1, weight=1)  # Стовпець для групи
        input_frame.columnconfigure(3, weight=0)  # Стовпець для номера пари
        input_frame.columnconfigure(5, weight=3)  # Стовпець для дисципліни (більший розмір)
        input_frame.columnconfigure(7, weight=1)  # Стовпець для аудиторії
        
        # Вибір групи (обов'язкове поле)
        ttk.Label(input_frame, text="Група *:", font=("Arial", 10)).grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.group_var = tk.StringVar()
        self.group_combo = ttk.Combobox(input_frame, textvariable=self.group_var, width=12, font=("Arial", 10))
        self.group_combo.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Номер пари
        ttk.Label(input_frame, text="№ пари *:", font=("Arial", 10)).grid(row=0, column=2, padx=5, pady=5, sticky=tk.W)
        self.lesson_var = tk.StringVar()
        ttk.Entry(input_frame, textvariable=self.lesson_var, width=8, font=("Arial", 10)).grid(row=0, column=3, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Дисципліна з автозаповненням
        ttk.Label(input_frame, text="Дисципліна:", font=("Arial", 10)).grid(row=0, column=4, padx=5, pady=5, sticky=tk.W)
        self.discipline_var = tk.StringVar()
        self.discipline_entry = ttk.Entry(input_frame, textvariable=self.discipline_var, width=35, font=("Arial", 10))
        self.discipline_entry.grid(row=0, column=5, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Створюємо Listbox для автозаповнення дисциплін
        self.discipline_listbox = tk.Listbox(input_frame, width=35, height=5, font=("Arial", 9))
        self.discipline_listbox.grid(row=1, column=5, padx=5, sticky=tk.W+tk.E)
        self.discipline_listbox.grid_remove()  # Спочатку приховуємо
        
        # Прив'язуємо події для автозаповнення
        self.discipline_var.trace("w", self.update_discipline_list)
        self.discipline_listbox.bind("<<ListboxSelect>>", self.on_discipline_select)
        
        # Вибір аудиторії
        ttk.Label(input_frame, text="Аудиторія:", font=("Arial", 10)).grid(row=0, column=6, padx=5, pady=5, sticky=tk.W)
        self.audience_var = tk.StringVar()
        self.audience_combo = ttk.Combobox(input_frame, textvariable=self.audience_var, width=12, values=self.audiences, font=("Arial", 10))
        self.audience_combo.grid(row=0, column=7, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Створюємо окремий фрейм для кнопки
        button_frame = ttk.Frame(replacement_frame)
        button_frame.pack(pady=10)
        
        # Кнопка додавання заміни (розташована по центру)
        add_btn = tk.Button(button_frame, text="Додати заміну", 
                           bg="#2ecc71", fg="white", font=("Arial", 11, "bold"),
                           relief=tk.RAISED, borderwidth=2, padx=15, pady=6,
                           command=self.add_replacement)
        add_btn.pack()
        
        # Секція для перегляду доданих замін
        replacements_frame = ttk.LabelFrame(main_frame, text="Додані заміни")
        replacements_frame.pack(fill=tk.X, pady=10)  # Змінено з fill=tk.BOTH, expand=True на fill=tk.X
        replacements_frame.configure(padding=10)
        
        # Створюємо фрейм для таблиці та скролбара
        tree_frame = ttk.Frame(replacements_frame)
        tree_frame.pack(fill=tk.X)  # Змінено з fill=tk.BOTH, expand=True на fill=tk.X
        
        # Створюємо таблицю для відображення замін
        columns = ("group", "lesson", "discipline", "audience")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=4)  # Зменшено висоту з 6 до 4
        
        # Задаємо заголовки колонок
        self.tree.heading("group", text="Група")
        self.tree.heading("lesson", text="№ пари")
        self.tree.heading("discipline", text="Дисципліна")
        self.tree.heading("audience", text="Аудиторія")
        
        # Задаємо ширину колонок
        self.tree.column("group", width=100)
        self.tree.column("lesson", width=80)
        self.tree.column("discipline", width=250)
        self.tree.column("audience", width=100)
        
        # Розміщуємо таблицю
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Додаємо скролбар
        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # Додаємо підтримку прокручування колесиком миші для таблиці
        def _on_tree_mousewheel(event):
            # Перевіряємо, чи є що прокручувати в таблиці
            if len(self.tree.get_children()) > 4:  # Більше ніж висота таблиці
                self.tree.yview_scroll(int(-1*(event.delta/120)), "units")
                return "break"  # Зупиняємо подальше поширення події
        
        # Прив'язуємо прокручування до таблиці
        self.tree.bind("<MouseWheel>", _on_tree_mousewheel)
        
        # Фрейм для кнопок управління замінами
        buttons_frame = ttk.Frame(replacements_frame)
        buttons_frame.pack(fill=tk.X, pady=(10, 0))
        
        # Кнопка видалення заміни
        remove_btn = tk.Button(buttons_frame, text="Видалити заміну", 
                              bg="#e74c3c", fg="white", font=("Arial", 10, "bold"),
                              relief=tk.RAISED, borderwidth=2, padx=10, pady=4,
                              command=self.remove_replacement)
        remove_btn.pack(side=tk.LEFT, padx=5)
        
        # Лічильник замін
        self.replacements_count_label = ttk.Label(buttons_frame, text="Замін: 0", font=("Arial", 10, "bold"))
        self.replacements_count_label.pack(side=tk.RIGHT, padx=5)
        
        # Секція практики
        practice_frame = ttk.LabelFrame(main_frame, text="Практика")
        practice_frame.pack(fill=tk.X, pady=10)
        practice_frame.configure(padding=10)  # Додаємо внутрішні відступи
        
        # Застосовуємо той самий стиль для заголовка фрейму практики
        # Стиль вже був змінений вище для всіх заголовків фреймів
        
        # Навчальна практика
        ttk.Label(practice_frame, text="Навчальна практика:", font=("Arial", 10)).grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.edu_practice_var = tk.StringVar()
        ttk.Entry(practice_frame, textvariable=self.edu_practice_var, width=40, font=("Arial", 10)).grid(
            row=0, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Тип практики
        ttk.Label(practice_frame, text="Тип практики:", font=("Arial", 10)).grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        
        # Радіокнопки для вибору типу практики
        radio_frame = ttk.Frame(practice_frame)
        radio_frame.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        self.practice_type = tk.StringVar(value="Виробнича")
        
        # Стилізовані радіокнопки
        rb1 = tk.Radiobutton(radio_frame, text="Виробнича", variable=self.practice_type, value="Виробнича", 
                           font=("Arial", 10), bg="#f0f0f0", selectcolor="#4a6984", indicatoron=0,
                           width=12, borderwidth=2, relief=tk.RAISED)
        rb1.pack(side=tk.LEFT, padx=5)
        
        rb2 = tk.Radiobutton(radio_frame, text="Переддипломна", variable=self.practice_type, value="Переддипломна", 
                           font=("Arial", 10), bg="#f0f0f0", selectcolor="#4a6984", indicatoron=0,
                           width=12, borderwidth=2, relief=tk.RAISED)
        rb2.pack(side=tk.LEFT, padx=5)
        
        # Інформація про практику
        ttk.Label(practice_frame, text="Інформація про практику:", font=("Arial", 10)).grid(row=2, column=0, padx=5, pady=5, sticky=tk.W)
        self.practice_info_var = tk.StringVar()
        ttk.Entry(practice_frame, textvariable=self.practice_info_var, width=40, font=("Arial", 10)).grid(
            row=2, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        
        # Секція суботніх пар
        saturday_frame = ttk.LabelFrame(main_frame, text="Налаштування суботніх пар")
        saturday_frame.pack(fill=tk.X, pady=10)
        saturday_frame.configure(padding=10)
        
        # Інформація про поточний день для суботи
        info_frame = ttk.Frame(saturday_frame)
        info_frame.pack(fill=tk.X, pady=5)
        
        self.saturday_info_label = ttk.Label(info_frame, text="", font=("Arial", 10, "italic"))
        self.saturday_info_label.pack(side=tk.LEFT)
        
        # Поле для базової дати
        base_date_frame = ttk.Frame(saturday_frame)
        base_date_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(base_date_frame, text="Базова дата (перший понеділок циклу):", font=("Arial", 10)).pack(side=tk.LEFT, padx=5)
        self.base_date_var = tk.StringVar()
        
        # Автоматично встановлюємо базову дату
        self.set_default_base_date()
        
        base_date_entry = ttk.Entry(base_date_frame, textvariable=self.base_date_var, width=15, font=("Arial", 10))
        base_date_entry.pack(side=tk.LEFT, padx=5)
        base_date_entry.bind('<KeyRelease>', self.on_base_date_change)
        
        # Кнопка для оновлення інформації про суботу
        update_saturday_btn = tk.Button(base_date_frame, text="Оновити інформацію", 
                                      bg="#3498db", fg="white", font=("Arial", 9),
                                      relief=tk.RAISED, borderwidth=1, padx=8, pady=4,
                                      command=self.update_saturday_info)
        update_saturday_btn.pack(side=tk.LEFT, padx=10)
        
        # Оновлюємо інформацію про суботу при ініціалізації
        self.update_saturday_info()
        
        # Фрейм для кнопок
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(pady=15)
        
        # Кнопка оновлення даних
        refresh_btn = tk.Button(buttons_frame, text="Оновити дані", 
                               bg="#e67e22", fg="white", font=("Arial", 11, "bold"),
                               relief=tk.RAISED, borderwidth=2, padx=12, pady=6,
                               command=self.refresh_data)
        refresh_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Кнопка формування бланку
        generate_btn = tk.Button(buttons_frame, text="Сформувати бланк", 
                               bg="#1a365d", fg="white", font=("Arial", 12, "bold"),
                               relief=tk.RAISED, borderwidth=2, padx=15, pady=8,
                               command=self.generate_form)
        generate_btn.pack(side=tk.LEFT)
        
        # Прив'язуємо прокручування колесиком до всіх віджетів
        self.bind_mousewheel_to_widget(main_frame)
    
    def validate_group_name(self, group_name):
        """Базова перевірка назви групи"""
        if not group_name or not group_name.strip():
            return False, "Назва групи не може бути порожньою"
        
        # Дозволяємо будь-яку назву групи, що не є порожньою
        return True, "Група прийнята"
    
    def validate_lesson_number(self, lesson_str):
        """Перевірка номера пари (може бути одна пара, кілька пар або діапазон)"""
        if not lesson_str or not lesson_str.strip():
            return False, "Номер пари не може бути порожнім"
        
        lesson_str = lesson_str.strip()
        
        # Перевіряємо різні формати
        
        # Формат: одна пара (1, 2, 3, ...)
        if re.match(r'^\d+$', lesson_str):
            lesson_num = int(lesson_str)
            if 1 <= lesson_num <= 8:  # Зазвичай пари від 1 до 8
                return True, f"Пара {lesson_num}"
            else:
                return False, f"Номер пари повинен бути від 1 до 8, введено: {lesson_num}"
        
        # Формат: діапазон (2-3, 1-4, ...)
        if re.match(r'^\d+-\d+$', lesson_str):
            start, end = map(int, lesson_str.split('-'))
            if 1 <= start <= 8 and 1 <= end <= 8:
                if start <= end:
                    return True, f"Пари {start}-{end}"
                else:
                    return False, f"Початок діапазону ({start}) не може бути більшим за кінець ({end})"
            else:
                return False, f"Номери пар у діапазоні повинні бути від 1 до 8, введено: {lesson_str}"
        
        # Формат: кілька пар через кому (1,3,5 або 1, 3, 5)
        if re.match(r'^\d+(\s*,\s*\d+)+$', lesson_str):
            lessons = [int(x.strip()) for x in lesson_str.split(',')]
            invalid_lessons = [l for l in lessons if not (1 <= l <= 8)]
            if invalid_lessons:
                return False, f"Недійсні номери пар: {invalid_lessons}. Номери пар повинні бути від 1 до 8"
            
            # Перевіряємо на дублікати
            if len(lessons) != len(set(lessons)):
                duplicates = [l for l in set(lessons) if lessons.count(l) > 1]
                return False, f"Знайдено дублікати номерів пар: {duplicates}"
            
            return True, f"Пари {', '.join(map(str, sorted(lessons)))}"
        
        # Формат: кілька пар через пробіл (1 3 5)
        if re.match(r'^\d+(\s+\d+)+$', lesson_str):
            lessons = [int(x) for x in lesson_str.split()]
            invalid_lessons = [l for l in lessons if not (1 <= l <= 8)]
            if invalid_lessons:
                return False, f"Недійсні номери пар: {invalid_lessons}. Номери пар повинні бути від 1 до 8"
            
            # Перевіряємо на дублікати
            if len(lessons) != len(set(lessons)):
                duplicates = [l for l in set(lessons) if lessons.count(l) > 1]
                return False, f"Знайдено дублікати номерів пар: {duplicates}"
            
            return True, f"Пари {', '.join(map(str, sorted(lessons)))}"
        
        return False, f"Неправильний формат номера пари: '{lesson_str}'. Дозволені формати: '1', '2-4', '1,3,5' або '1 3 5'"
    
    def get_saturday_schedule_day(self, target_date, base_date=None):
        """
        Визначає, за який день тижня проводяться суботні пари для заданої дати.
        
        Args:
            target_date (datetime.date): Дата суботи, для якої потрібно визначити день
            base_date (datetime.date): Базова дата початку циклу (перший понеділок)
        
        Returns:
            str: Назва дня тижня українською мовою
        """
        if base_date is None:
            # Якщо базова дата не задана, використовуємо початок навчального року
            # Припускаємо, що навчальний рік починається 1 вересня
            current_year = target_date.year
            if target_date.month < 9:  # Якщо до вересня, то попередній навчальний рік
                current_year -= 1
            base_date = datetime.date(current_year, 9, 1)
            
            # Знаходимо перший понеділок після 1 вересня
            while base_date.weekday() != 0:  # 0 = понеділок
                base_date += datetime.timedelta(days=1)
        
        # Обчислюємо кількість тижнів від базової дати до цільової дати
        days_diff = (target_date - base_date).days
        weeks_passed = days_diff // 7
        
        # Цикл: понеділок(0) → вівторок(1) → середа(2) → четвер(3) → п'ятниця(4) → знову понеділок
        day_cycle = weeks_passed % 5
        
        days_ua = {
            0: "понеділок",
            1: "вівторок", 
            2: "середа",
            3: "четвер",
            4: "п'ятниця"
        }
        
        return days_ua[day_cycle]
    
    def is_saturday(self, date_str=None):
        """Перевіряє, чи є задана дата суботою"""
        try:
            # Перевіряємо день тижня з self.weekday (окрема змінна)
            if hasattr(self, 'weekday') and self.weekday:
                return self.weekday.lower() == "субота"
            return False
        except:
            return False
    
    def _parse_date_from_text(self, date_text):
        """Парсить дату з тексту формату 'день місяць, день_тижня'"""
        try:
            date_parts = date_text.split()
            if len(date_parts) >= 2:
                day = int(date_parts[0])
                month_name = date_parts[1].rstrip(',')
                
                if month_name in MONTHS_UA:
                    month = MONTHS_UA[month_name]
                    current_year = datetime.date.today().year
                    return datetime.date(current_year, month, day)
            return None
        except:
            return None
    
    def set_default_base_date(self):
        """Встановлює базову дату за замовчуванням"""
        try:
            # Отримуємо поточну дату з форми
            if hasattr(self, 'date_text') and self.date_text:
                target_date = self._parse_date_from_text(self.date_text)
                if target_date:
                    # Знаходимо початок навчального року
                    if target_date.month < 9:
                        academic_year_start = target_date.year - 1
                    else:
                        academic_year_start = target_date.year
                    
                    base_date = datetime.date(academic_year_start, 9, 1)
                    # Знаходимо перший понеділок
                    while base_date.weekday() != 0:
                        base_date += datetime.timedelta(days=1)
                    
                    self.base_date_var.set(base_date.strftime("%d.%m.%Y"))
                    return
            
            # Якщо не вдалося парсити дату, використовуємо поточний навчальний рік
            today = datetime.date.today()
            if today.month < 9:
                academic_year_start = today.year - 1
            else:
                academic_year_start = today.year
            
            base_date = datetime.date(academic_year_start, 9, 1)
            while base_date.weekday() != 0:
                base_date += datetime.timedelta(days=1)
            
            self.base_date_var.set(base_date.strftime("%d.%m.%Y"))
        except:
            # У випадку помилки встановлюємо дату за замовчуванням
            self.base_date_var.set("01.09.2024")
    
    def on_base_date_change(self, event=None):
        """Обробник зміни базової дати"""
        self.update_saturday_info()
    
    def update_saturday_info(self):
        """Оновлює інформацію про суботні пари"""
        try:
            if not hasattr(self, 'date_text') or not self.date_text:
                self.saturday_info_label.config(text="Дата не задана")
                return
            
            # Перевіряємо, чи це субота
            if not self.is_saturday(self.date_text):
                self.saturday_info_label.config(text="Цей день не є суботою")
                return
            
            # Парсимо дату з тексту
            target_date = self._parse_date_from_text(self.date_text)
            if target_date:
                # Парсимо базову дату
                base_date_str = self.base_date_var.get()
                if base_date_str:
                    base_parts = base_date_str.split('.')
                    if len(base_parts) == 3:
                        base_date = datetime.date(int(base_parts[2]), int(base_parts[1]), int(base_parts[0]))
                        
                        # Визначаємо день тижня для суботи
                        schedule_day = self.get_saturday_schedule_day(target_date, base_date)
                        
                        self.saturday_info_label.config(
                            text=f"Суботні пари проводяться за розкладом: {schedule_day}",
                            foreground="#2c3e50"
                        )
                        return
            
            self.saturday_info_label.config(text="Помилка при обробці дати")
        except Exception as e:
            self.saturday_info_label.config(text=f"Помилка: {str(e)}")
    
    def on_department_change(self, event=None):
        dept = self.dept_var.get()
        if dept in self.groups:
            self.group_combo['values'] = self.groups[dept]
        else:
            self.group_combo['values'] = []
    
    def update_discipline_list(self, *args):
        typed = self.discipline_var.get().lower()
        print(f"\n\nВведено текст: '{typed}'")
        print(f"Кількість дисциплін у списку: {len(self.disciplines)}")
        print(f"Перші 5 дисциплін: {self.disciplines[:5] if len(self.disciplines) >= 5 else self.disciplines}")
        
        if typed == '':
            self.discipline_listbox.grid_remove()
            print("Список приховано, бо введений текст порожній")
        else:
            # Позиціонуємо список під полем вводу
            x, y, width, height = self.discipline_entry.winfo_x(), self.discipline_entry.winfo_y(), self.discipline_entry.winfo_width(), self.discipline_entry.winfo_height()
            
            # Встановлюємо список точно під полем вводу
            self.discipline_listbox.grid(row=1, column=5, padx=10, sticky=tk.W+tk.E)
            self.discipline_listbox.delete(0, tk.END)
            
            # Збільшуємо ширину списку, щоб він був не менший за поле вводу
            self.discipline_listbox.config(width=max(40, width // 8))  # Ширина не менше 40 символів
            
            matching_disciplines = []
            for discipline in self.disciplines:
                if typed in discipline.lower():
                    matching_disciplines.append(discipline)
                    self.discipline_listbox.insert(tk.END, discipline)
            
            # Якщо знайдено дисципліни, відображаємо список
            if matching_disciplines:
                self.discipline_listbox.grid()
                # Підлаштовуємо висоту списку в залежності від кількості знайдених елементів
                height = min(10, len(matching_disciplines))  # Максимум 10 елементів
                self.discipline_listbox.config(height=height)
            else:
                self.discipline_listbox.grid_remove()
            
            print(f"Знайдено {len(matching_disciplines)} дисциплін, що відповідають запиту")
            if matching_disciplines:
                print(f"Знайдені дисципліни: {matching_disciplines}")
    
    def on_discipline_select(self, event):
        if self.discipline_listbox.curselection():
            selected = self.discipline_listbox.get(self.discipline_listbox.curselection())
            self.discipline_var.set(selected)
            self.discipline_listbox.grid_remove()
    
    def refresh_data(self):
        """Оновлення даних з бази даних без перезапуску програми"""
        try:
            # Показуємо повідомлення про початок оновлення
            messagebox.showinfo("Оновлення даних", "Оновлення даних з бази даних...")
            
            # Оновлюємо дані в базі даних
            if self.db.refresh_data():
                # Очищаємо локальний кеш
                self._clear_data_cache()
                
                # Зберігаємо старі дані для порівняння
                old_departments = list(self.departments) if self._departments else []
                old_groups = dict(self.groups) if self._groups else {}
                
                # Оновлюємо список аудиторій
                old_audiences = self.audiences.copy()
                self.audiences = self.db.get_audiences()
                
                # Оновлюємо список дисциплін
                old_disciplines = self.disciplines.copy()
                self.disciplines = self.db.get_disciplines()
                
                # Оновлюємо віджети форми
                self.update_form_widgets()
                
                # Показуємо інформацію про зміни
                changes = []
                if len(self.departments) != len(old_departments):
                    changes.append(f"Відділення: {len(old_departments)} → {len(self.departments)}")
                
                total_old_groups = sum(len(groups) for groups in old_groups.values())
                total_new_groups = sum(len(groups) for groups in self.groups.values())
                if total_new_groups != total_old_groups:
                    changes.append(f"Групи: {total_old_groups} → {total_new_groups}")
                
                if len(self.audiences) != len(old_audiences):
                    changes.append(f"Аудиторії: {len(old_audiences)} → {len(self.audiences)}")
                
                if len(self.disciplines) != len(old_disciplines):
                    changes.append(f"Дисципліни: {len(old_disciplines)} → {len(self.disciplines)}")
                
                if changes:
                    change_text = "Виявлено зміни:\n" + "\n".join(changes)
                else:
                    change_text = "Структура даних не змінилася"
                
                messagebox.showinfo("Оновлення завершено", 
                                  f"Дані успішно оновлено з бази даних!\n\n{change_text}")
            else:
                messagebox.showwarning("Помилка оновлення", 
                                     "Не вдалося оновити дані з бази даних.\n"
                                     "Перевірте підключення до бази даних.")
        
        except Exception as e:
            messagebox.showerror("Помилка", f"Помилка при оновленні даних: {e}")
    
    def update_form_widgets(self):
        """Оновлення віджетів форми після оновлення даних"""
        try:
            # Оновлюємо список відділень
            current_dept = self.dept_var.get()
            self.dept_combo['values'] = self.departments
            
            # Якщо поточне відділення більше не існує, очищуємо вибір
            if current_dept not in self.departments:
                self.dept_var.set("")
                self.group_combo['values'] = []
                self.group_var.set("")
            else:
                # Оновлюємо список груп для поточного відділення
                self.on_department_change()
            
            # Оновлюємо список аудиторій
            current_audience = self.audience_var.get()
            self.audience_combo['values'] = self.audiences
            
            # Якщо поточна аудиторія більше не існує, очищуємо вибір
            if current_audience not in self.audiences:
                self.audience_var.set("")
            
            print("Віджети форми успішно оновлено")
            
        except Exception as e:
            print(f"Помилка при оновленні віджетів форми: {e}")
    
    def add_replacement(self):
        group = self.group_var.get().strip()
        lesson = self.lesson_var.get().strip()
        discipline = self.discipline_var.get().strip()
        audience = self.audience_var.get().strip()
        
        # Перевіряємо обов'язкові поля
        if not group or not lesson:
            messagebox.showwarning("Попередження", "Заповніть обов'язкові поля: Група та № пари")
            return
        
        # Валідація групи
        group_valid, group_message = self.validate_group_name(group)
        if not group_valid:
            messagebox.showerror("Помилка валідації групи", group_message)
            return
        
        # Валідація номера пари
        lesson_valid, lesson_message = self.validate_lesson_number(lesson)
        if not lesson_valid:
            messagebox.showerror("Помилка валідації номера пари", lesson_message)
            return
        
        # Add to replacements list
        self.replacements.append({
            "group": group,
            "lesson": lesson,
            "discipline": discipline,
            "audience": audience
        })
        
        # Add to treeview
        self.tree.insert("", tk.END, values=(group, lesson, discipline, audience))
        
        # Clear entry fields
        self.group_var.set("")
        self.lesson_var.set("")
        self.discipline_var.set("")
        self.audience_var.set("")
        
        # Оновлюємо відображення списку замін
        self.update_replacements_display()
    
    def update_replacements_display(self):
        """Оновлює відображення кількості замін"""
        count = len(self.replacements)
        self.replacements_count_label.config(text=f"Замін: {count}")
        
        # Якщо є заміни, підсвічуємо лічильник
        if count > 0:
            self.replacements_count_label.config(foreground="#2ecc71")  # Зелений колір
        else:
            self.replacements_count_label.config(foreground="black")  # Чорний колір
    
    def remove_replacement(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Попередження", "Виберіть заміну для видалення")
            return
        
        # Get index of selected item
        index = self.tree.index(selected[0])
        
        # Remove from replacements list
        self.replacements.pop(index)
        
        # Remove from treeview
        self.tree.delete(selected[0])
        
        # Оновлюємо відображення списку замін
        self.update_replacements_display()
    
    def generate_form(self):
        # Використовуємо дані, які були передані при створенні форми
        date_text = self.date_text
        weekday = self.weekday
        week_type = self.week_type
        duty_group = self.duty_group_text
        duty_teacher = self.duty_teacher_text
        dorm_teacher = self.dorm_teacher_text
        
        # Отримуємо дані про практику
        edu_practice = self.edu_practice_var.get().strip()
        practice_type_text = self.practice_type.get()
        practice_info = self.practice_info_var.get().strip()
        
        # Якщо немає реальних замін, повідомляємо користувача
        if not self.replacements:
            messagebox.showinfo("Інформація", "Додайте хоча б одну заміну для формування бланку")
            return
        
        try:
            # Створюємо документ Word
            doc = Document()
            
            # Встановлюємо шрифт для всього документа
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            
            # Налаштовуємо поля сторінки
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)
            
            # Додаємо заголовок "ЗАТВЕРДЖУЮ"
            header_table = doc.add_table(rows=1, cols=2)
            header_table.style = 'Normal Table'
            
            for row in header_table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = tcPr.first_child_found_in("w:tcBorders")
                    if tcBorders:
                        tcBorders.getparent().remove(tcBorders)
            
            # Права частина (ЗАТВЕРДЖУЮ)
            right_cell = header_table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_run = right_para.add_run("ЗАТВЕРДЖУЮ")
            right_run.bold = True
            right_run.font.size = Pt(8)
            right_para.paragraph_format.space_after = Pt(0)
            right_para.paragraph_format.line_spacing = 1.0
            
            position_para = right_cell.add_paragraph()
            position_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            position_run = position_para.add_run("Заступник директора з навчальної роботи")
            position_run.font.size = Pt(8)
            position_para.paragraph_format.space_after = Pt(0)
            position_para.paragraph_format.line_spacing = 1.0
            
            initials_para = right_cell.add_paragraph()
            initials_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            initials_run = initials_para.add_run("Балдич Л. В.")
            initials_run.font.size = Pt(8)
            initials_para.paragraph_format.space_after = Pt(0)
            initials_para.paragraph_format.line_spacing = 1.0
            
            # Додаємо заголовок документа
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("Зміни до розкладу занять")
            title_run.bold = True
            title_run.font.size = Pt(15)
            title_para.paragraph_format.space_after = Pt(3)
            title_para.paragraph_format.line_spacing = 1.0
            
            # Додаємо дату
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_parts = date_text.split()
            
            # Формуємо текст дати
            if len(date_parts) >= 2:
                short_date = f"{date_parts[0]} {date_parts[1]}"
            else:
                short_date = date_text
            
            # Перевіряємо чи це субота
            if self.is_saturday():
                # Визначаємо за яким днем тижня проводяться пари
                target_date = self._parse_date_from_text(date_text)
                if target_date:
                    # Парсимо базову дату
                    base_date_str = self.base_date_var.get()
                    if base_date_str:
                        base_parts = base_date_str.split('.')
                        if len(base_parts) == 3:
                            base_date = datetime.date(int(base_parts[2]), int(base_parts[1]), int(base_parts[0]))
                            schedule_day = self.get_saturday_schedule_day(target_date, base_date)
                            # Форматуємо текст для суботи
                            date_run = date_para.add_run(f"на {short_date}, суботу, за {schedule_day}\nзаняття за {week_type.lower()}")
                        else:
                            # Якщо не вдалося парсити базову дату, використовуємо стандартний формат
                            date_run = date_para.add_run(f"на {short_date}, {weekday}\nнавчання за {week_type.lower()}")
                    else:
                        # Якщо базова дата не встановлена, використовуємо стандартний формат
                        date_run = date_para.add_run(f"на {short_date}, {weekday}\nнавчання за {week_type.lower()}")
                else:
                    # Якщо не вдалося парсити дату, використовуємо стандартний формат
                    date_run = date_para.add_run(f"на {short_date}, {weekday}\nнавчання за {week_type.lower()}")
            else:
                # Для звичайних днів тижня (не субота)
                date_run = date_para.add_run(f"на {short_date}, {weekday}\nнавчання за {week_type.lower()}")
            
            date_run.bold = True
            date_run.font.size = Pt(15)
            date_para.paragraph_format.space_after = Pt(3)
            date_para.paragraph_format.line_spacing = 1.0
            
            # Отримуємо структуру відділень
            dept_structure = self.db.get_department_structure()
            departments = [dept["name"] for dept in dept_structure]
            
            if not departments:
                messagebox.showinfo("Інформація", "Немає відділень у базі даних")
                return
            
            # Розподіляємо заміни по відділеннях
            dept_replacements = {dept: [] for dept in departments}
            
            for replacement in self.replacements:
                group = replacement["group"]
                for dept, groups in self.groups.items():
                    if group in groups:
                        if dept in dept_replacements:
                            dept_replacements[dept].append(replacement)
                        break
            
            # Визначаємо структуру
            general_edu_name = "Загальноосвітньої підготовки"
            has_general_edu = any(general_edu_name in dept["name"] for dept in dept_structure)
            
            ordered_departments = []
            general_edu_dept = None
            
            for dept_info in sorted(dept_structure, key=lambda x: x["order"]):
                if general_edu_name in dept_info["name"]:
                    general_edu_dept = dept_info["name"]
                else:
                    ordered_departments.append(dept_info["name"])
            
            # Розраховуємо кількість рядків для таблиці
            total_rows = 0
            
            # Рядки для Загальноосвітньої підготовки
            if has_general_edu and general_edu_dept:
                total_rows += 1  # Заголовок
                gen_edu_repls = dept_replacements.get(general_edu_dept, [])
                total_rows += max(1, len(gen_edu_repls))  # Мінімум 1 порожній рядок
            
            # Рядки для інших відділень (по 2 в рядку)
            num_departments = len(ordered_departments)
            num_rows = (num_departments + 1) // 2
            
            for row_idx in range(num_rows):
                total_rows += 1  # Заголовки
                
                # Рахуємо максимум замін в цьому рядку
                max_repls = 0
                for col_idx in range(2):
                    dept_idx = row_idx * 2 + col_idx
                    if dept_idx < num_departments:
                        dept_name = ordered_departments[dept_idx]
                        dept_repls = dept_replacements.get(dept_name, [])
                        max_repls = max(max_repls, len(dept_repls))
                
                total_rows += max(1, max_repls)
            
            # Рядки для практик та чергових
            total_rows += 5  # 2 для практик + 3 для чергових
            
            # Створюємо таблицю
            main_table = doc.add_table(rows=total_rows, cols=6)
            main_table.style = 'Table Grid'
            main_table.autofit = False
            main_table.allow_autofit = False
            
            # Встановлюємо ширину колонок
            from docx.oxml.ns import nsdecls
            tbl = main_table._element
            tblPr = tbl.tblPr
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>')
            tblPr.append(tblW)
            
            tblGrid = tbl.tblGrid
            if tblGrid is not None:
                tbl.remove(tblGrid)
            
            col_widths = [
                int(1.5 * 567), int(1.0 * 567), int(6.5 * 567),
                int(1.5 * 567), int(1.0 * 567), int(6.5 * 567),
            ]
            
            tblGrid_xml = '<w:tblGrid %s>' % nsdecls('w')
            for width in col_widths:
                tblGrid_xml += f'<w:gridCol w:w="{width}"/>'
            tblGrid_xml += '</w:tblGrid>'
            
            new_tblGrid = parse_xml(tblGrid_xml)
            tbl.insert(1, new_tblGrid)
            
            # Видаляємо всі межі та встановлюємо ширину
            for row in main_table.rows:
                for col_idx, cell in enumerate(row.cells):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    tcBorders = tcPr.first_child_found_in("w:tcBorders")
                    if tcBorders is not None:
                        tcBorders.getparent().remove(tcBorders)
                    
                    width = col_widths[col_idx]
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width}" w:type="dxa"/>')
                    tcPr.append(tcW)
            
            # Функції для роботи з таблицею
            def add_borders(cell, top=False, bottom=False, left=False, right=False):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                borders_xml = '<w:tcBorders %s>' % nsdecls('w')
                if top:
                    borders_xml += '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                else:
                    borders_xml += '<w:top w:val="nil"/>'
                if bottom:
                    borders_xml += '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                else:
                    borders_xml += '<w:bottom w:val="nil"/>'
                if left:
                    borders_xml += '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                else:
                    borders_xml += '<w:left w:val="nil"/>'
                if right:
                    borders_xml += '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                else:
                    borders_xml += '<w:right w:val="nil"/>'
                borders_xml += '</w:tcBorders>'
                
                tcBorders = parse_xml(borders_xml)
                tcPr.append(tcBorders)
            
            def add_text_to_cell(cell, text, bold=False, gray_fill=False, align_left=False):
                cell.text = text
                paragraph = cell.paragraphs[0]
                if align_left:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if paragraph.runs:
                    run = paragraph.runs[0]
                    if bold:
                        run.bold = True
                    run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                if gray_fill:
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Заповнюємо таблицю
            current_row = 0
            
            # Загальноосвітня підготовка
            if has_general_edu and general_edu_dept:
                header_cell = main_table.cell(current_row, 0)
                for col in range(1, 6):
                    header_cell.merge(main_table.cell(current_row, col))
                
                add_text_to_cell(header_cell, general_edu_dept, bold=True, gray_fill=True)
                add_borders(header_cell, top=True, bottom=True, left=True, right=True)
                current_row += 1
                
                # Заміни для Загальноосвітньої підготовки
                gen_edu_repls = dept_replacements.get(general_edu_dept, [])
                
                # Групуємо за групами
                grouped = {}
                for repl in gen_edu_repls:
                    group = repl["group"]
                    if group not in grouped:
                        grouped[group] = []
                    grouped[group].append(repl)
                
                if grouped:
                    # Розподіляємо групи по колонках (ліва та права частина)
                    sorted_groups = sorted(grouped.keys())
                    groups_left = []
                    groups_right = []
                    
                    for idx, group_name in enumerate(sorted_groups):
                        if idx % 2 == 0:
                            groups_left.append(group_name)
                        else:
                            groups_right.append(group_name)
                    
                    # Розгортаємо заміни для лівої частини
                    left_repls = []
                    for group_name in groups_left:
                        for idx, repl in enumerate(grouped[group_name]):
                            left_repls.append({
                                'group': group_name if idx == 0 else '',
                                'lesson': repl['lesson'],
                                'discipline': repl.get('discipline', ''),
                                'audience': repl.get('audience', '')
                            })
                    
                    # Розгортаємо заміни для правої частини
                    right_repls = []
                    for group_name in groups_right:
                        for idx, repl in enumerate(grouped[group_name]):
                            right_repls.append({
                                'group': group_name if idx == 0 else '',
                                'lesson': repl['lesson'],
                                'discipline': repl.get('discipline', ''),
                                'audience': repl.get('audience', '')
                            })
                    
                    # Заповнюємо рядки
                    max_rows = max(len(left_repls), len(right_repls))
                    
                    for row_idx in range(max_rows):
                        # Ліва частина
                        if row_idx < len(left_repls):
                            repl = left_repls[row_idx]
                            
                            # Група
                            group_cell = main_table.cell(current_row, 0)
                            add_text_to_cell(group_cell, repl['group'])
                            add_borders(group_cell, left=True)
                            
                            # Номер пари
                            lesson_cell = main_table.cell(current_row, 1)
                            add_text_to_cell(lesson_cell, repl['lesson'])
                            
                            # Предмет
                            text_cell = main_table.cell(current_row, 2)
                            discipline = repl['discipline']
                            audience = repl['audience'].strip()
                            if discipline and audience:
                                text = f"{discipline}, ауд. {audience}"
                            elif discipline:
                                text = discipline
                            elif audience:
                                text = f"ауд. {audience}"
                            else:
                                text = ""
                            add_text_to_cell(text_cell, text, align_left=True)
                            add_borders(text_cell, right=True)
                        else:
                            # Порожні комірки зліва
                            for col in range(3):
                                cell = main_table.cell(current_row, col)
                                add_text_to_cell(cell, "")
                                if col == 0:
                                    add_borders(cell, left=True)
                                elif col == 2:
                                    add_borders(cell, right=True)
                        
                        # Права частина
                        if row_idx < len(right_repls):
                            repl = right_repls[row_idx]
                            
                            # Група
                            group_cell = main_table.cell(current_row, 3)
                            add_text_to_cell(group_cell, repl['group'])
                            add_borders(group_cell, left=True)
                            
                            # Номер пари
                            lesson_cell = main_table.cell(current_row, 4)
                            add_text_to_cell(lesson_cell, repl['lesson'])
                            
                            # Предмет
                            text_cell = main_table.cell(current_row, 5)
                            discipline = repl['discipline']
                            audience = repl['audience'].strip()
                            if discipline and audience:
                                text = f"{discipline}, ауд. {audience}"
                            elif discipline:
                                text = discipline
                            elif audience:
                                text = f"ауд. {audience}"
                            else:
                                text = ""
                            add_text_to_cell(text_cell, text, align_left=True)
                            add_borders(text_cell, right=True)
                        else:
                            # Порожні комірки справа
                            for col in range(3, 6):
                                cell = main_table.cell(current_row, col)
                                add_text_to_cell(cell, "")
                                if col == 3:
                                    add_borders(cell, left=True)
                                elif col == 5:
                                    add_borders(cell, right=True)
                        
                        current_row += 1
                else:
                    # Порожній рядок
                    for col in range(6):
                        cell = main_table.cell(current_row, col)
                        add_text_to_cell(cell, "")
                        borders = {}
                        if col == 0:
                            borders['left'] = True
                        if col == 5:
                            borders['right'] = True
                        if col == 2:
                            borders['right'] = True
                        if col == 3:
                            borders['left'] = True
                        add_borders(cell, **borders)
                    current_row += 1
                
                # Нижня межа
                for col in range(6):
                    cell = main_table.cell(current_row - 1, col)
                    add_borders(cell, bottom=True, left=(col==0 or col==3), right=(col==2 or col==5))
            
            # Інші відділення
            for row_idx in range(num_rows):
                # Заголовки
                for col_idx in range(2):
                    dept_idx = row_idx * 2 + col_idx
                    if dept_idx < num_departments:
                        dept_name = ordered_departments[dept_idx]
                        
                        start_col = col_idx * 3
                        header_cell = main_table.cell(current_row, start_col)
                        header_cell.merge(main_table.cell(current_row, start_col + 1))
                        header_cell.merge(main_table.cell(current_row, start_col + 2))
                        
                        add_text_to_cell(header_cell, dept_name, bold=True, gray_fill=True)
                        add_borders(header_cell, top=True, bottom=True, left=True, right=True)
                
                current_row += 1
                
                # Заміни
                max_repls = 0
                dept_repls_list = [[], []]
                
                for col_idx in range(2):
                    dept_idx = row_idx * 2 + col_idx
                    if dept_idx < num_departments:
                        dept_name = ordered_departments[dept_idx]
                        repls = dept_replacements.get(dept_name, [])
                        
                        # Групуємо
                        grouped = {}
                        for repl in repls:
                            group = repl["group"]
                            if group not in grouped:
                                grouped[group] = []
                            grouped[group].append(repl)
                        
                        # Розгортаємо в список
                        flat_list = []
                        for group_name in sorted(grouped.keys()):
                            for idx, repl in enumerate(grouped[group_name]):
                                flat_list.append({
                                    'group': group_name if idx == 0 else '',
                                    'lesson': repl['lesson'],
                                    'discipline': repl.get('discipline', ''),
                                    'audience': repl.get('audience', '')
                                })
                        
                        dept_repls_list[col_idx] = flat_list
                        max_repls = max(max_repls, len(flat_list))
                
                # Заповнюємо рядки
                for repl_idx in range(max(1, max_repls)):
                    for col_idx in range(2):
                        start_col = col_idx * 3
                        
                        if repl_idx < len(dept_repls_list[col_idx]):
                            repl = dept_repls_list[col_idx][repl_idx]
                            
                            # Група
                            group_cell = main_table.cell(current_row, start_col)
                            add_text_to_cell(group_cell, repl['group'])
                            add_borders(group_cell, left=True, right=True)
                            
                            # Номер пари
                            lesson_cell = main_table.cell(current_row, start_col + 1)
                            add_text_to_cell(lesson_cell, repl['lesson'])
                            add_borders(lesson_cell, right=True)
                            
                            # Предмет
                            text_cell = main_table.cell(current_row, start_col + 2)
                            discipline = repl['discipline']
                            audience = repl['audience'].strip()
                            if discipline and audience:
                                text = f"{discipline}, ауд. {audience}"
                            elif discipline:
                                text = discipline
                            elif audience:
                                text = f"ауд. {audience}"
                            else:
                                text = ""
                            add_text_to_cell(text_cell, text, align_left=True)
                            add_borders(text_cell, right=True)
                        else:
                            # Порожні комірки
                            for sub_col in range(3):
                                cell = main_table.cell(current_row, start_col + sub_col)
                                add_text_to_cell(cell, "")
                                if sub_col == 0:
                                    add_borders(cell, left=True, right=True)
                                elif sub_col == 1:
                                    add_borders(cell, right=True)
                                elif sub_col == 2:
                                    add_borders(cell, right=True)
                    
                    current_row += 1
                
                # Нижня межа
                for col in range(6):
                    cell = main_table.cell(current_row - 1, col)
                    add_borders(cell, bottom=True, left=(col==0 or col==3), right=(col==2 or col==5))
            
            # Практики
            practice_cell1 = main_table.cell(current_row, 0)
            practice_cell1.merge(main_table.cell(current_row, 1))
            practice_cell1.merge(main_table.cell(current_row, 2))
            add_text_to_cell(practice_cell1, "НАВЧАЛЬНА ПРАКТИКА", bold=True, gray_fill=True)
            add_borders(practice_cell1, top=True, bottom=True, left=True, right=True)
            
            practice_cell2 = main_table.cell(current_row, 3)
            practice_cell2.merge(main_table.cell(current_row, 4))
            practice_cell2.merge(main_table.cell(current_row, 5))
            add_text_to_cell(practice_cell2, "ВИРОБНИЧА ПРАКТИКА", bold=True, gray_fill=True)
            add_borders(practice_cell2, top=True, bottom=True, left=True, right=True)
            
            current_row += 1
            
            practice_info1 = main_table.cell(current_row, 0)
            practice_info1.merge(main_table.cell(current_row, 1))
            practice_info1.merge(main_table.cell(current_row, 2))
            add_text_to_cell(practice_info1, "--------")
            add_borders(practice_info1, top=True, bottom=True, left=True, right=True)
            
            practice_info2 = main_table.cell(current_row, 3)
            practice_info2.merge(main_table.cell(current_row, 4))
            practice_info2.merge(main_table.cell(current_row, 5))
            add_text_to_cell(practice_info2, "--------")
            add_borders(practice_info2, top=True, bottom=True, left=True, right=True)
            
            current_row += 1
            
            # Чергові
            duty_cell = main_table.cell(current_row, 0)
            for col in range(1, 6):
                duty_cell.merge(main_table.cell(current_row, col))
            add_text_to_cell(duty_cell, f"Чергова група: {duty_group}", bold=False, align_left=True)
            add_borders(duty_cell, top=True, bottom=True, left=True, right=True)
            
            current_row += 1
            
            teacher_cell = main_table.cell(current_row, 0)
            for col in range(1, 6):
                teacher_cell.merge(main_table.cell(current_row, col))
            add_text_to_cell(teacher_cell, f"Черговий викладач: {duty_teacher}", bold=False, align_left=True)
            add_borders(teacher_cell, top=True, bottom=True, left=True, right=True)
            
            current_row += 1
            
            dorm_cell = main_table.cell(current_row, 0)
            for col in range(1, 6):
                dorm_cell.merge(main_table.cell(current_row, col))
            add_text_to_cell(dorm_cell, f"Черговий викладач у гуртожитках: {dorm_teacher}", bold=False, align_left=True)
            add_borders(dorm_cell, top=True, bottom=True, left=True, right=True)
            
            # Формуємо назву файлу
            date_parts = date_text.split()
            if len(date_parts) >= 2:
                day = date_parts[0]
                month = date_parts[1]
                default_filename = f"{day} {month}.docx"
            else:
                default_filename = f"{date_text}.docx"
            
            # Зберігаємо файл
            if self.replacements_dir and self.academic_year:
                month_folders = {
                    "січня": "01-Січень", "лютого": "02-Лютий", "березня": "03-Березень",
                    "квітня": "04-Квітень", "травня": "05-Травень", "червня": "06-Червень",
                    "липня": "07-Липень", "серпня": "08-Серпень", "вересня": "09-Вересень",
                    "жовтня": "10-Жовтень", "листопада": "11-Листопад", "грудня": "12-Грудень"
                }
                
                if len(date_parts) >= 2:
                    month_folder = month_folders.get(date_parts[1], "")
                    year_dir = os.path.join(self.replacements_dir, self.academic_year)
                    if not os.path.exists(year_dir):
                        os.makedirs(year_dir)
                    
                    month_dir = os.path.join(year_dir, month_folder)
                    if month_folder and not os.path.exists(month_dir):
                        os.makedirs(month_dir)
                    
                    if month_folder:
                        full_path = os.path.join(month_dir, default_filename)
                    else:
                        full_path = os.path.join(year_dir, default_filename)
                    
                    doc.save(full_path)
                    messagebox.showinfo("Успіх", f"Бланк замін збережено у файл {os.path.basename(full_path)}\nШлях: {full_path}")
                else:
                    year_dir = os.path.join(self.replacements_dir, self.academic_year)
                    if not os.path.exists(year_dir):
                        os.makedirs(year_dir)
                    
                    full_path = os.path.join(year_dir, default_filename)
                    doc.save(full_path)
                    messagebox.showinfo("Успіх", f"Бланк замін збережено у файл {os.path.basename(full_path)}\nШлях: {full_path}")
            else:
                filename = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
                    initialfile=default_filename
                )
                
                if filename:
                    doc.save(filename)
                    messagebox.showinfo("Успіх", f"Бланк замін збережено у файл {os.path.basename(filename)}")
                else:
                    messagebox.showinfo("Інформація", "Збереження файлу скасовано")
        
        except Exception as e:
            messagebox.showerror("Помилка", f"Помилка при створенні документу: {str(e)}")
            import traceback
            traceback.print_exc()

            # Формуємо назву файлу у форматі "день місяць"
            # Витягуємо день і місяць з дати
            date_parts = date_text.split()
            if len(date_parts) >= 2:
                day = date_parts[0]
                month = date_parts[1]
                default_filename = f"{day} {month}.docx"
            else:
                # Якщо не вдалося розібрати дату, використовуємо всю дату
                default_filename = f"{date_text}.docx"
            
            # Перевіряємо, чи передано шлях для збереження
            if self.replacements_dir and self.academic_year:
                # Словник для перетворення назви місяця в папку
                month_folders = {
                    "січня": "01-Січень",
                    "лютого": "02-Лютий",
                    "березня": "03-Березень",
                    "квітня": "04-Квітень",
                    "травня": "05-Травень",
                    "червня": "06-Червень",
                    "липня": "07-Липень",
                    "серпня": "08-Серпень",
                    "вересня": "09-Вересень",
                    "жовтня": "10-Жовтень",
                    "листопада": "11-Листопад",
                    "грудня": "12-Грудень"
                }
                
                if len(date_parts) >= 2:
                    month_folder = month_folders.get(date_parts[1], "")
                    
                    # Створюємо шлях для збереження
                    year_dir = os.path.join(self.replacements_dir, self.academic_year)
                    if not os.path.exists(year_dir):
                        os.makedirs(year_dir)
                    
                    # Створюємо папку місяця, якщо вона не існує
                    month_dir = os.path.join(year_dir, month_folder)
                    if month_folder and not os.path.exists(month_dir):
                        os.makedirs(month_dir)
                    
                    # Формуємо повний шлях до файлу
                    if month_folder:
                        full_path = os.path.join(month_dir, default_filename)
                    else:
                        full_path = os.path.join(year_dir, default_filename)
                    
                    # Зберігаємо файл без діалогу
                    doc.save(full_path)
                    messagebox.showinfo("Успіх", f"Бланк замін збережено у файл {os.path.basename(full_path)}\nШлях: {full_path}")
                else:
                    # Якщо не вдалося визначити місяць, зберігаємо в папку року
                    year_dir = os.path.join(self.replacements_dir, self.academic_year)
                    if not os.path.exists(year_dir):
                        os.makedirs(year_dir)
                    
                    full_path = os.path.join(year_dir, default_filename)
                    doc.save(full_path)
                    messagebox.showinfo("Успіх", f"Бланк замін збережено у файл {os.path.basename(full_path)}\nШлях: {full_path}")
            else:
                # Якщо шлях не передано, використовуємо стандартний діалог збереження
                filename = filedialog.asksaveasfilename(
                    initialdir=os.path.expanduser("~") + "/Desktop",
                    initialfile=default_filename,
                    defaultextension=".docx",
                    filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
                )
                
                if filename:  # Якщо користувач не скасував діалог
                    doc.save(filename)
                    messagebox.showinfo("Успіх", f"Бланк замін збережено у файл {os.path.basename(filename)}")
                else:
                    messagebox.showinfo("Інформація", "Збереження файлу скасовано")
        
        except Exception as e:
            messagebox.showerror("Помилка", f"Помилка при створенні документу: {e}")
